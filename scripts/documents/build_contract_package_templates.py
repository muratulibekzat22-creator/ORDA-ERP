"""Build the approved ALTYN SAPA contract template.

Design basis: the `contract_negotiation_brief` document preset with a compact
`customer_pack` header, overridden to strict A4 contract geometry. The legal
clauses are copied verbatim from the approved V1 template; the customer memo is
intentionally excluded and generated as a separate PDF by ORDA.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "resources" / "documents" / "templates" / "contract-altyn-sapa-v2.docx"
INK = "202124"
MUTED = "5F6368"
GOLD = "B68A3A"
GOLD_SOFT = "F6EFE2"
LINE = "D8D3CA"


def set_cell_shading(cell, fill: str):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge, attributes in edges.items():
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in attributes.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    props = row._tr.get_or_add_trPr()
    props.append(OxmlElement("w:cantSplit"))


def format_run(run, size=9.5, bold=False, color=INK, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def compact_paragraph(paragraph, before=0, after=1.5, line=1.0, keep=False):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_with_next = keep


def add_text(document, value, *, bold=False, size=9.5, color=INK, before=0, after=1.5, keep=False, align=None):
    paragraph = document.add_paragraph()
    compact_paragraph(paragraph, before=before, after=after, keep=keep)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(value)
    format_run(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(document, value):
    paragraph = add_text(document, value, bold=True, size=10, color=INK, before=3, after=1.5, keep=True)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_field(paragraph, instruction: str):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.extend([begin, instruction_node, separate, text, end])


def write_cell(cell, label: str, value: str, *, accent=False):
    cell.text = ""
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    compact_paragraph(paragraph, after=0)
    label_run = paragraph.add_run(f"{label}\n")
    format_run(label_run, size=7.6, bold=True, color=GOLD if accent else MUTED)
    value_run = paragraph.add_run(value)
    format_run(value_run, size=9.5, bold=accent, color=INK)


def project_table(document):
    table = document.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.15)
    table.columns[1].width = Cm(9.15)
    fields = [
        ("КЛИЕНТ", "{{clientFullName}}", "ИИН / ТЕЛЕФОН", "{{clientIin}} · {{clientPhone}}"),
        ("ЗАКАЗ", "№ {{orderNumber}}", "АДРЕС МОНТАЖА", "{{installationAddress}}"),
        ("МАТЕРИАЛ", "{{stairMaterial}}", "КАРКАС", "{{frameType}} {{frameComment}}"),
        ("ОГРАЖДЕНИЕ / БАЛЯСИНА", "{{balusterType}}", "СТОЙКА / ЦВЕТ", "{{supportType}} · {{color}}"),
        ("ПОДСВЕТКА", "{{lightingText}}", "ОБШИВКА", "{{claddingText}}"),
        ("ДОСТАВКА / МОНТАЖ", "{{deliveryText}} · {{installationText}}", "ДОПОЛНИТЕЛЬНО", "{{additionalDetails}}"),
    ]
    border = {"val": "single", "sz": "4", "color": LINE}
    for left_label, left_value, right_label, right_value in fields:
        row = table.add_row()
        prevent_row_split(row)
        write_cell(row.cells[0], left_label, left_value)
        write_cell(row.cells[1], right_label, right_value)
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)
    return table


def payment_table(document):
    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.15)
    table.columns[1].width = Cm(9.15)
    write_cell(table.cell(0, 0), "ОБЩАЯ СТОИМОСТЬ", "{{contractAmount}} ₸", accent=True)
    write_cell(table.cell(0, 1), "СРОК", "{{termCalendarDays}} календарных дней")
    write_cell(table.cell(1, 0), "ПЕРВЫЙ ПЛАТЁЖ", "{{paymentSchedulePrimary}}")
    write_cell(table.cell(1, 1), "ОСТАТОК / ГРАФИК", "{{paymentScheduleBalance}}")
    border = {"val": "single", "sz": "5", "color": GOLD}
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_shading(cell, GOLD_SOFT)
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)


def requisites_table(document):
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.15)
    table.columns[1].width = Cm(9.15)
    row = table.rows[0]
    prevent_row_split(row)
    left = (
        "КЛИЕНТ\n\n"
        "ФИО: {{clientFullName}}\n"
        "ИИН: {{clientIin}}\n"
        "Телефон: {{clientPhone}}\n"
        "Адрес: {{clientAddress}}\n\n"
        "Подпись: ____________________"
    )
    right = (
        "ИСПОЛНИТЕЛЬ\n\n"
        "ТОО «{{companyName}}»\n"
        "БИН: {{companyBin}}\n"
        "ИИК: {{companyIik}}\n"
        "Банк: {{companyBank}}\n"
        "БИК: {{companyBik}}\n"
        "Телефон / WhatsApp:\n{{companyPhoneLines}}\n"
        "Адрес: {{companyAddress}}\n"
        "Директор: {{directorFullName}}\n\n"
        "Подпись: ____________________"
    )
    for cell, value in zip(row.cells, (left, right)):
        cell.text = ""
        set_cell_margins(cell, top=100, start=110, bottom=100, end=110)
        set_cell_shading(cell, "FAFAF9")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        paragraph = cell.paragraphs[0]
        compact_paragraph(paragraph, after=0, line=0.95)
        run = paragraph.add_run(value)
        format_run(run, size=9.5)
        border = {"val": "single", "sz": "4", "color": LINE}
        set_cell_border(cell, top=border, bottom=border, start=border, end=border)


def header_and_footer(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.45)

    header = section.header
    paragraph = header.paragraphs[0]
    compact_paragraph(paragraph, after=0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("ALTYN SAPA COMPANY")
    format_run(run, size=8, bold=True, color=GOLD)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    compact_paragraph(paragraph, after=0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Страница ")
    format_run(run, size=8, color=MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" из ")
    format_run(run, size=8, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def build():
    document = Document()
    header_and_footer(document)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)

    title = add_text(document, "ALTYN SAPA COMPANY", bold=True, size=10, color=GOLD, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.keep_with_next = True
    add_text(document, "ДОГОВОР № {{contractNumber}}", bold=True, size=17, after=0, keep=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(document, "изготовления, поставки и монтажа лестницы", size=9.5, color=MUTED, after=3, keep=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    date_table = document.add_table(rows=1, cols=2)
    date_table.autofit = False
    date_table.columns[0].width = Cm(9.15)
    date_table.columns[1].width = Cm(9.15)
    write_cell(date_table.cell(0, 0), "ГОРОД", "г. {{contractCity}}")
    write_cell(date_table.cell(0, 1), "ДАТА И ВРЕМЯ", "«{{contractDay}}» {{contractMonth}} {{contractYear}} г., {{contractTime}}")
    prevent_row_split(date_table.rows[0])
    for cell in date_table.rows[0].cells:
        set_cell_shading(cell, "FAFAF9")
        border = {"val": "single", "sz": "4", "color": LINE}
        set_cell_border(cell, top=border, bottom=border, start=border, end=border)

    add_text(
        document,
        "{{clientFullName}}, ИИН {{clientIin}}, именуемый(ая) в дальнейшем «Клиент», с одной стороны, и ТОО «ALTYN SAPA COMPANY», БИН 220540017969, в лице директора {{directorFullName}}, действующего на основании Устава, именуемое в дальнейшем «Исполнитель», с другой стороны, совместно именуемые «Стороны», заключили настоящий Договор о нижеследующем:",
        after=2,
    )
    project_table(document)

    add_heading(document, "1. ПРЕДМЕТ ДОГОВОРА")
    add_text(document, "1.1. Исполнитель обязуется изготовить, поставить Клиенту и выполнить монтаж лестницы (далее — «Лестница») в соответствии с согласованными параметрами заказа, замером и условиями настоящего Договора, а Клиент обязуется принять результат работ и оплатить его.")
    add_text(document, "1.2. Материал лестницы: {{stairMaterial}}.")
    add_text(document, "1.3. Тип ограждения / балясины: {{balusterType}}.")
    add_text(document, "1.4. Адрес монтажа: {{installationAddress}}.")

    add_heading(document, "2. СТОИМОСТЬ ДОГОВОРА И ПОРЯДОК ОПЛАТЫ")
    payment_table(document)
    add_text(document, "2.1. Общая стоимость Договора составляет {{contractAmount}} ({{contractAmountWords}}) тенге. Статус НДС указывается в соответствии с действующим налоговым статусом Исполнителя.")
    add_text(document, "2.2. В стоимость включаются согласованные Сторонами материалы, изготовление, доставка и монтаж в объеме, предусмотренном заказом.")
    add_text(document, "2.3. Оплата производится банковским переводом на расчётный счёт Исполнителя либо иным согласованным и допустимым способом.")
    add_text(document, "2.4. Согласованный график оплаты:", keep=True)
    add_text(document, "Первый платёж — {{prepaymentPercent}}% от стоимости Договора, что составляет {{prepaymentAmount}} ({{prepaymentAmountWords}}) тенге — {{prepaymentDueText}}.")
    add_text(document, "Оставшаяся сумма — {{balancePercent}}% от стоимости Договора, что составляет {{balanceAmount}} ({{balanceAmountWords}}) тенге — {{balanceDueText}}.")
    add_text(document, "2.5. При согласовании 100% оплаты ORDA формирует отдельную редакцию пункта 2.4: «Оплата производится в размере 100% стоимости Договора — {{contractAmount}} ({{contractAmountWords}}) тенге — {{fullPaymentDueText}}».")
    add_text(document, "2.6. При рассрочке через банк или иную финансовую организацию порядок оплаты определяется условиями одобренного финансирования и подтверждённым графиком платежей.")
    add_text(document, "2.7. Изменение согласованного графика оплаты оформляется по соглашению Сторон и фиксируется в документах ORDA.")

    page_break = document.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    add_heading(document, "3. ОБЯЗАННОСТИ СТОРОН")
    add_text(document, "3.1. Клиент обязуется обеспечить доступ на объект, предоставить необходимые исходные данные, принять выполненные работы и произвести оплату в согласованные сроки.")
    add_text(document, "3.2. Исполнитель обязуется выполнить согласованный объём работ в соответствии с параметрами заказа и условиями настоящего Договора.")
    add_text(document, "3.3. Права и обязанности Сторон, не урегулированные настоящим Договором, определяются законодательством Республики Казахстан.")

    add_heading(document, "4. СРОКИ И УСЛОВИЯ ВЫПОЛНЕНИЯ")
    add_text(document, "4.1. Срок изготовления и выполнения согласованного объёма работ: {{termCalendarDays}} календарных дней с даты {{termStartCondition}}.")
    add_text(document, "4.2. Плановая дата готовности / завершения: {{plannedCompletionDate}}.")
    add_text(document, "4.3. Клиент обеспечивает готовность объекта к доставке и монтажу, включая свободный доступ, подготовленную площадку и необходимые коммуникации.")
    add_text(document, "4.4. При возникновении обстоятельств, препятствующих выполнению работ и зависящих от Клиента, сроки могут быть перенесены на период устранения таких обстоятельств с соответствующей фиксацией.")

    add_heading(document, "5. ГАРАНТИЯ")
    add_text(document, "5.1. Гарантия на Лестницу: {{warrantyText}} с даты подписания акта приёмки-передачи / фактической приёмки результата работ.")
    add_text(document, "5.2. Гарантия не распространяется на повреждения, возникшие вследствие неправильной эксплуатации, механического воздействия, пожара, затопления, самостоятельной переделки либо иных действий, не связанных с качеством изготовления и монтажа Исполнителем.")

    add_heading(document, "6. ФОРС-МАЖОР И ПРОЧИЕ УСЛОВИЯ")
    add_text(document, "6.1. При возникновении обстоятельств непреодолимой силы Сторона, для которой возникла невозможность исполнения обязательств, уведомляет другую Сторону в разумный срок.")
    add_text(document, "6.2. Стороны руководствуются законодательством Республики Казахстан по вопросам, не урегулированным настоящим Договором.")
    add_text(document, "6.3. Изменения и дополнения к настоящему Договору оформляются в письменной форме либо иным способом, позволяющим достоверно установить согласованную волю Сторон в соответствии с применимым законодательством.")
    add_text(document, "6.4. Договор составлен в двух экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из Сторон.")

    add_heading(document, "7. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН")
    requisites_table(document)

    document.core_properties.title = "Договор ALTYN SAPA COMPANY"
    document.core_properties.subject = "Изготовление, поставка и монтаж лестницы"
    document.core_properties.author = "ALTYN SAPA COMPANY"
    document.settings.element.append(OxmlElement("w:updateFields"))
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
